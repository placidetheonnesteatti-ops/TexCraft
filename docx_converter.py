# -*- coding: utf-8 -*-
"""
Convertisseur Word (.docx) -> LaTeX (.tex)

Gère :
 - titres (Heading 1/2/3 -> section/subsection/subsubsection)
 - gras / italique / souligné
 - listes à puces et numérotées
 - images (extraites et incluses avec \\includegraphics, centrées)
 - tableaux simples
 - correction orthographique/grammaticale optionnelle
"""

import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from text_utils import escape_latex


def _iter_block_items(parent):
    """
    Parcourt les paragraphes ET les tableaux d'un document dans leur ordre
    réel d'apparition (au lieu de traiter tous les paragraphes puis tous
    les tableaux séparément, ce que fait doc.paragraphs / doc.tables).
    Fonctionne aussi à l'intérieur des cellules de tableau (tableaux imbriqués).
    """
    if hasattr(parent, "element"):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


HEADING_MAP = {
    "Heading 1": "section",
    "Titre 1": "section",
    "Heading 2": "subsection",
    "Titre 2": "subsection",
    "Heading 3": "subsubsection",
    "Titre 3": "subsubsection",
}

def _get_page_geometry(doc: Document) -> dict:
    """Récupère la taille de page et les marges réelles du document Word."""
    section = doc.sections[0]
    return {
        "paperwidth": section.page_width.cm if section.page_width else 21.0,
        "paperheight": section.page_height.cm if section.page_height else 29.7,
        "top": section.top_margin.cm if section.top_margin else 2.5,
        "bottom": section.bottom_margin.cm if section.bottom_margin else 2.5,
        "left": section.left_margin.cm if section.left_margin else 2.5,
        "right": section.right_margin.cm if section.right_margin else 2.5,
    }


def _get_body_font_size_pt(doc: Document) -> float:
    """Récupère la taille de police du style 'Normal' du document Word (par défaut 11pt)."""
    try:
        size = doc.styles["Normal"].font.size
        if size is not None:
            return round(size.pt, 1)
    except Exception:
        pass
    return 11.0


def _build_preamble(doc: Document) -> str:
    geo = _get_page_geometry(doc)
    font_pt = _get_body_font_size_pt(doc)
    leading_pt = round(font_pt * 1.2, 1)
    # Classe de base la plus proche (10/11/12pt) — la taille exacte est
    # ensuite appliquée par-dessus via \fontsize pour coller au point près
    # à la taille du document Word d'origine.
    base = min((10, 11, 12), key=lambda b: abs(b - font_pt))

    return r"""\documentclass[a4paper,%dpt]{article}
\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
\usepackage[french]{babel}
\usepackage{graphicx}
\usepackage{geometry}
\usepackage{enumitem}
\geometry{paperwidth=%.3fcm,paperheight=%.3fcm,top=%.3fcm,bottom=%.3fcm,left=%.3fcm,right=%.3fcm}
\renewcommand{\normalsize}{\fontsize{%.1fpt}{%.1fpt}\selectfont}
\normalsize

\begin{document}

""" % (base, geo["paperwidth"], geo["paperheight"], geo["top"], geo["bottom"],
       geo["left"], geo["right"], font_pt, leading_pt)


LATEX_END = "\n\\end{document}\n"


def _run_to_latex(run) -> str:
    text = escape_latex(run.text)
    if not text:
        return ""
    if run.bold:
        text = r"\textbf{%s}" % text
    if run.italic:
        text = r"\emph{%s}" % text
    if run.underline:
        text = r"\underline{%s}" % text
    return text


def _paragraph_to_latex(paragraph) -> str:
    return "".join(_run_to_latex(r) for r in paragraph.runs) or escape_latex(paragraph.text)


def _extract_images(doc: Document, images_dir: str) -> dict:
    """Extrait toutes les images du document vers images_dir. Retourne rel_id -> chemin fichier."""
    os.makedirs(images_dir, exist_ok=True)
    rel_to_path = {}
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            image_part = rel.target_part
            ext = image_part.content_type.split("/")[-1]
            ext = "jpg" if ext == "jpeg" else ext
            filename = f"image_{len(rel_to_path)+1}.{ext}"
            path = os.path.join(images_dir, filename)
            with open(path, "wb") as f:
                f.write(image_part.blob)
            rel_to_path[rel_id] = filename
    return rel_to_path


def _paragraph_has_image(paragraph, rel_to_path):
    """Renvoie le nom de fichier image si le paragraphe contient une image inline."""
    blips = paragraph._p.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    )
    for blip in blips:
        rel_id = blip.get(qn("r:embed"))
        if rel_id in rel_to_path:
            return rel_to_path[rel_id]
    return None


def _table_to_latex(table) -> str:
    ncols = len(table.columns)
    col_spec = "|" + "l|" * ncols
    lines = [r"\begin{table}[h]", r"\centering", r"\begin{tabular}{%s}" % col_spec, r"\hline"]
    for row in table.rows:
        cells = [escape_latex(cell.text).replace("\n", " ") for cell in row.cells]
        lines.append(" & ".join(cells) + r" \\ \hline")
    lines += [r"\end{tabular}", r"\end{table}", ""]
    return "\n".join(lines)


def convert_docx_to_latex(input_path: str, output_dir: str, title: str = "",
                           checker=None, progress_cb=None) -> str:
    """
    Convertit un fichier .docx en .tex.
    Retourne le chemin du fichier .tex généré.
    """
    os.makedirs(output_dir, exist_ok=True)
    images_dir = os.path.join(output_dir, "images")
    doc = Document(input_path)

    if not title:
        title = os.path.splitext(os.path.basename(input_path))[0]

    rel_to_path = _extract_images(doc, images_dir)

    body_parts = [_build_preamble(doc)]

    in_itemize = False
    in_enumerate = False

    def close_lists():
        nonlocal in_itemize, in_enumerate
        if in_itemize:
            body_parts.append(r"\end{itemize}" + "\n")
            in_itemize = False
        if in_enumerate:
            body_parts.append(r"\end{enumerate}" + "\n")
            in_enumerate = False

    block_items = list(_iter_block_items(doc))
    total = len(block_items)

    for i, block in enumerate(block_items):
        if progress_cb:
            progress_cb(i + 1, total, "Traitement du document...")

        # --- Tableau : inséré exactement à sa position réelle ---
        if isinstance(block, Table):
            close_lists()
            body_parts.append(_table_to_latex(block))
            continue

        # --- Paragraphe ---
        para = block
        style_name = para.style.name if para.style else ""

        # Image dans le paragraphe (insérée à sa position réelle)
        img_file = _paragraph_has_image(para, rel_to_path)
        if img_file:
            close_lists()
            body_parts.append(
                "\\begin{figure}[h]\n\\centering\n"
                f"\\includegraphics[width=0.8\\textwidth]{{images/{img_file}}}\n"
                "\\end{figure}\n"
            )
            # Le paragraphe peut aussi contenir du texte en plus de l'image
            text_raw = para.text.strip()
            if not text_raw:
                continue

        text_raw = para.text.strip()
        if not text_raw:
            continue

        # Titres
        if style_name in HEADING_MAP:
            close_lists()
            content = escape_latex(text_raw)
            # Étoile = titre NON numéroté (comme dans Word, qui ne numérote
            # pas automatiquement les titres, contrairement à LaTeX par défaut)
            body_parts.append(f"\\{HEADING_MAP[style_name]}*{{{content}}}\n")
            continue

        # Listes (détection par nom de style ; python-docx ne fiabilise pas
        # toujours numPr selon l'origine du document)
        style_lower = style_name.lower()
        is_numbered = "list number" in style_lower or "liste numérotée" in style_lower or "numbered" in style_lower
        is_bullet = (not is_numbered) and (
            "list bullet" in style_lower or "liste à puces" in style_lower or "bullet" in style_lower
        )

        content = _paragraph_to_latex(para)
        if checker is not None:
            content = checker.correct(content)

        if is_numbered:
            if not in_enumerate:
                close_lists()
                body_parts.append(r"\begin{enumerate}")
                in_enumerate = True
            body_parts.append(f"  \\item {content}")
        elif is_bullet:
            if not in_itemize:
                close_lists()
                body_parts.append(r"\begin{itemize}")
                in_itemize = True
            body_parts.append(f"  \\item {content}")
        else:
            close_lists()
            body_parts.append(content + "\n")

    close_lists()
    body_parts.append(LATEX_END)

    tex_content = "\n".join(body_parts)
    out_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".tex")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(tex_content)

    return out_path
